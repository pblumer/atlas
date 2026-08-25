# -*- coding: utf-8 -*-
"""Fill the BIT P042-Hi01 template with the answers from isds-konzept.md.

Keeps the template's cover page, change-control tables, headers/footers and
styles; replaces the body from chapter 1 onwards with the converted markdown.
Needs python-docx; see README.md next to this file.
"""
import re
import sys

from docx import Document
from docx.shared import Pt

if len(sys.argv) != 4:
    raise SystemExit('usage: render.py <isds-konzept.md> <P042-Hi01-template.docx> <out.docx>')
MD, TEMPLATE, OUT = sys.argv[1], sys.argv[2], sys.argv[3]

doc = Document(TEMPLATE)
body = doc.element.body

# --- 1. cover page fields -------------------------------------------------
COVER = {
    'Klassifizierung': 'intern',
    'Status': 'in Arbeit',
    'Version': '0.1 (Entwurf Hersteller/Projektteam, Basis Template V4.4 / BIT-Template V120)',
    'Datum': '25.08.2026',
    'Autor/Autoren': 'Atlas-Projektteam',
}
for tbl in doc.tables:
    for row in tbl.rows:
        if len(row.cells) == 2:
            key = row.cells[0].text.strip()
            if key in COVER:
                cell = row.cells[1]
                cell.paragraphs[0].text = ''
                for p in cell.paragraphs[1:]:
                    p._element.getparent().remove(p._element)
                cell.paragraphs[0].add_run(COVER.pop(key))
for p in doc.paragraphs:
    if p.text.strip() == '<Projektname / Schutzobjektname>':
        for r in p.runs[1:]:
            r._element.getparent().remove(r._element)
        p.runs[0].text = 'Atlas — BPMN 2.x Workflow Engine'

# --- 2. cut the template body from chapter 1 on ----------------------------
start = None
for p in doc.paragraphs:
    if p.style.name in ('Heading 1', 'berschrift1') or p.style.name.startswith('Heading'):
        if p.text.strip() == 'Generelle Anmerkungen':
            start = p._element
            break
if start is None:
    raise SystemExit('anchor paragraph "Generelle Anmerkungen" not found')
el = start
while el is not None:
    nxt = el.getnext()
    if el.tag.endswith('}sectPr'):
        break
    body.remove(el)
    el = nxt

sect = body.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr')

def add(el):
    """Append a block element before the final sectPr."""
    if sect is not None:
        sect.addprevious(el)
    else:
        body.append(el)

def new_par(text='', style=None):
    p = doc.add_paragraph(style=style)
    body.remove(p._element)
    add(p._element)
    if text:
        write_inline(p, text)
    return p

INLINE = re.compile(r'(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))')

def write_inline(par, text):
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            par.add_run(part[2:-2]).bold = True
        elif part.startswith('`') and part.endswith('`'):
            r = par.add_run(part[1:-1]); r.font.name = 'Consolas'; r.font.size = Pt(9)
        elif part.startswith('['):
            m = re.match(r'\[([^\]]+)\]\(([^)]+)\)', part)
            par.add_run(m.group(1))
        else:
            par.add_run(part)

def add_table(rows):
    cols = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=cols)
    t.style = 'TableGrid'
    body.remove(t._element)
    add(t._element)
    for i, row in enumerate(rows):
        for j in range(cols):
            cell = t.cell(i, j)
            par = cell.paragraphs[0]
            par.style = doc.styles['Normal']
            txt = row[j] if j < len(row) else ''
            write_inline(par, txt)
            for run in par.runs:
                run.font.size = Pt(8)
                if i == 0:
                    run.bold = True
    return t

# --- 3. convert the markdown ----------------------------------------------
lines = open(MD, encoding='utf-8').read().splitlines()
i = 0
# skip the front matter: start at chapter 1
while i < len(lines) and not lines[i].startswith('# 1 '):
    i += 1

HEAD = re.compile(r'^(#{1,3})\s+(?:\d+(?:\.\d+)*\s+)?(.*)$')
CELL = re.compile(r'^\s*\|(.+)\|\s*$')

while i < len(lines):
    line = lines[i]
    if line.strip() in ('---', '***', '___'):
        i += 1; continue
    if not line.strip():
        i += 1; continue
    m = HEAD.match(line)
    if m:
        level = len(m.group(1))
        new_par(m.group(2), style=f'Heading {level}')
        i += 1; continue
    if line.startswith('```'):
        fence = line[3:].strip()
        i += 1
        buf = []
        while i < len(lines) and not lines[i].startswith('```'):
            buf.append(lines[i]); i += 1
        i += 1
        if fence == 'mermaid':
            new_par('⟨Architekturskizze hier als Bild einfügen. Der Quelltext des Diagramms '
                    '(Mermaid) folgt, damit die Skizze reproduzierbar bleibt.⟩')
        for src in buf:
            cp = new_par()
            cp.paragraph_format.space_after = Pt(0)
            r = cp.add_run(src.replace('\t', '    '))
            r.font.name = 'Consolas'; r.font.size = Pt(8)
        new_par()
        continue
    if CELL.match(line):
        rows = []
        while i < len(lines) and CELL.match(lines[i]):
            cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
            if not all(re.fullmatch(r':?-{2,}:?', c) for c in cells):
                rows.append(cells)
            i += 1
        add_table(rows)
        new_par()
        continue
    if re.match(r'^\s*[-*]\s+', line):
        while i < len(lines) and (re.match(r'^\s*[-*]\s+', lines[i]) or (lines[i].startswith('  ') and lines[i].strip())):
            if re.match(r'^\s*[-*]\s+', lines[i]):
                indent = len(lines[i]) - len(lines[i].lstrip())
                style = 'List Bullet 2' if indent >= 2 else 'List Bullet'
                try:
                    p = new_par(re.sub(r'^\s*[-*]\s+', '', lines[i]), style=style)
                except KeyError:
                    p = new_par(re.sub(r'^\s*[-*]\s+', '', lines[i]), style='List Paragraph')
            else:
                write_inline(p, ' ' + lines[i].strip())
            i += 1
        continue
    if re.match(r'^\s*\d+\.\s+', line):
        while i < len(lines) and (re.match(r'^\s*\d+\.\s+', lines[i]) or (lines[i].startswith('   ') and lines[i].strip())):
            if re.match(r'^\s*\d+\.\s+', lines[i]):
                try:
                    p = new_par(re.sub(r'^\s*\d+\.\s+', '', lines[i]), style='List Number')
                except KeyError:
                    p = new_par(re.sub(r'^\s*\d+\.\s+', '', lines[i]), style='List Paragraph')
            else:
                write_inline(p, ' ' + lines[i].strip())
            i += 1
        continue
    if line.startswith('>'):
        buf = []
        while i < len(lines) and lines[i].startswith('>'):
            buf.append(lines[i].lstrip('> ').rstrip()); i += 1
        p = new_par(' '.join(buf))
        for r in p.runs:
            r.italic = True
        continue
    buf = [line]
    i += 1
    while i < len(lines) and lines[i].strip() and not lines[i].startswith(('#', '|', '>', '```', '- ', '* ')) \
            and not re.match(r'^\s*\d+\.\s+', lines[i]) and lines[i].strip() not in ('---',):
        buf.append(lines[i].strip()); i += 1
    new_par(' '.join(buf))

doc.save(OUT)
print('written', OUT)
